"""
RAG Service - The heart of the system.
Handles: document ingestion → chunking → embeddings → ChromaDB → retrieval → Gemini LLM
"""
import os
import re
import base64
import asyncio
from typing import List, Dict, Optional, AsyncGenerator
from pathlib import Path

import google.generativeai as genai
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_chroma import Chroma
from langchain.schema import Document

from app.config import get_settings

import time
from google.api_core.exceptions import ResourceExhausted

settings = get_settings()

# Configure Gemini
genai.configure(api_key=settings.gemini_api_key)

# langchain-google-genai requires GOOGLE_API_KEY env var
os.environ.setdefault("GOOGLE_API_KEY", settings.gemini_api_key)


def generate_content_with_retry(model, contents, stream=False, max_retries=5, initial_delay=3.0):
    delay = initial_delay
    for attempt in range(max_retries):
        try:
            if stream:
                return model.generate_content(contents, stream=True)
            else:
                return model.generate_content(contents)
        except ResourceExhausted as e:
            if attempt == max_retries - 1:
                raise e
            print(f"[WARN] Gemini API 429 Rate Limit hit. Retrying in {delay}s... (Attempt {attempt + 1}/{max_retries})")
            time.sleep(delay)
            delay *= 2
        except Exception as e:
            raise e


# Embedding model
_embeddings = None


def get_embeddings():
    global _embeddings
    if _embeddings is None:
        _embeddings = GoogleGenerativeAIEmbeddings(
            model="models/gemini-embedding-2",
            google_api_key=settings.gemini_api_key,
        )
    return _embeddings


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT LOADING
# ══════════════════════════════════════════════════════════════════════════════

def load_pdf(file_path: str) -> List[Document]:
    """Extract text from PDF using pypdf with page metadata."""
    from pypdf import PdfReader
    reader = PdfReader(file_path)
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": os.path.basename(file_path), "page": i + 1}
            ))
    # If pypdf yields nothing (scanned PDF), use Gemini OCR
    if not docs:
        docs = ocr_pdf_with_gemini(file_path)
    return docs


def _detect_image_mime(img_data: bytes) -> str:
    """Sniff the MIME type of raw image bytes from magic numbers."""
    if img_data[:3] == b'\xff\xd8\xff':
        return "image/jpeg"
    if img_data[:4] == b'\x89PNG':
        return "image/png"
    if img_data[:4] in (b'GIF8', b'GIF9'):
        return "image/gif"
    if img_data[:4] == b'RIFF' and img_data[8:12] == b'WEBP':
        return "image/webp"
    return "image/png"  # safe fallback


def ocr_pdf_with_gemini(file_path: str) -> List[Document]:
    """Use Gemini multimodal to OCR scanned PDF pages."""
    from pypdf import PdfReader
    reader = PdfReader(file_path)
    model = genai.GenerativeModel("gemini-2.0-flash")
    docs = []
    for i, page in enumerate(reader.pages):
        try:
            images = page.images
            if images:
                img_data = images[0].data
                mime_type = _detect_image_mime(img_data)
                response = generate_content_with_retry(model, [
                    "Extract all text from this document image accurately:",
                    {"mime_type": mime_type, "data": base64.b64encode(img_data).decode()}
                ])
                text = response.text
                if text.strip():
                    docs.append(Document(
                        page_content=text,
                        metadata={"source": os.path.basename(file_path), "page": i + 1, "ocr": True}
                    ))
        except Exception:
            continue
    return docs


def load_docx(file_path: str) -> List[Document]:
    """Extract text from DOCX."""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [Document(
        page_content=text,
        metadata={"source": os.path.basename(file_path), "page": 1}
    )]


def load_txt(file_path: str) -> List[Document]:
    """Load plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [Document(
        page_content=text,
        metadata={"source": os.path.basename(file_path), "page": 1}
    )]


def load_document(file_path: str, file_type: str) -> List[Document]:
    ext = file_type.lower()
    if ext == "pdf":
        return load_pdf(file_path)
    elif ext in ("docx", "doc"):
        return load_docx(file_path)
    elif ext == "txt":
        return load_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ══════════════════════════════════════════════════════════════════════════════
#  CHUNKING
# ══════════════════════════════════════════════════════════════════════════════

def chunk_documents(docs: List[Document], chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )
    return splitter.split_documents(docs)


# ══════════════════════════════════════════════════════════════════════════════
#  VECTOR STORE
# ══════════════════════════════════════════════════════════════════════════════

def get_vectorstore(collection_name: str) -> Chroma:
    return Chroma(
        collection_name=collection_name,
        embedding_function=get_embeddings(),
        persist_directory=settings.chroma_db_path,
    )


def ingest_documents(file_path: str, file_type: str, collection_name: str) -> int:
    """Load → chunk → embed → store. Returns chunk count."""
    raw_docs = load_document(file_path, file_type)
    chunks = chunk_documents(raw_docs)
    if not chunks:
        return 0
    vectorstore = get_vectorstore(collection_name)
    vectorstore.add_documents(chunks)
    return len(chunks)


def delete_collection(collection_name: str):
    """Remove a document's collection from ChromaDB."""
    import chromadb
    client = chromadb.PersistentClient(path=settings.chroma_db_path)
    try:
        client.delete_collection(collection_name)
    except Exception:
        pass


# ══════════════════════════════════════════════════════════════════════════════
#  RETRIEVAL & QA
# ══════════════════════════════════════════════════════════════════════════════

def retrieve_context(query: str, collection_names: List[str], top_k: int = 6) -> List[Document]:
    """Search across multiple collections and return top-k combined results."""
    all_docs = []
    for name in collection_names:
        try:
            vs = get_vectorstore(name)
            results = vs.similarity_search(query, k=top_k)
            all_docs.extend(results)
        except Exception as e:
            print(f"[WARN] Similarity search failed for {name}: {e}. Trying raw retrieval...")
            try:
                res = vs.get()
                if res and "documents" in res:
                    for doc_text, meta in zip(res["documents"], res["metadatas"]):
                        all_docs.append(Document(page_content=doc_text, metadata=meta))
            except Exception:
                continue
    # Deduplicate and sort by relevance (already sorted per-collection)
    seen = set()
    unique = []
    for d in all_docs:
        key = d.page_content[:100]
        if key not in seen:
            seen.add(key)
            unique.append(d)
    return unique[:top_k]


def build_context_string(docs: List[Document]) -> tuple[str, List[Dict]]:
    """Build context block and source citations from retrieved docs."""
    context_parts = []
    sources = []
    for i, doc in enumerate(docs):
        src = doc.metadata.get("source", "Unknown")
        page = doc.metadata.get("page", "?")
        context_parts.append(f"[Source {i+1}: {src}, Page {page}]\n{doc.page_content}")
        sources.append({"source": src, "page": page, "snippet": doc.page_content[:200]})
    return "\n\n---\n\n".join(context_parts), sources


# ══════════════════════════════════════════════════════════════════════════════
#  LOCAL FALLBACK GENERATORS (WHEN GEMINI API IS RATE-LIMITED / OFFLINE)
# ══════════════════════════════════════════════════════════════════════════════

def generate_fallback_rag_answer(query: str, context: str) -> str:
    """Generate a clean, structured summary of retrieved context locally."""
    lines = [line.strip() for line in context.split('\n') if line.strip()]
    relevant_lines = []
    for line in lines:
        if any(keyword in line.lower() for keyword in query.lower().split() if len(keyword) > 2):
            relevant_lines.append(line)
    
    if not relevant_lines:
        relevant_lines = lines[:10]
        
    summary_bullets = []
    for line in relevant_lines:
        if not line.startswith('[Source'):
            cleaned = re.sub(r'^[-*\d.\s]+', '', line).strip()
            if cleaned and cleaned not in summary_bullets and len(cleaned) > 10:
                summary_bullets.append(cleaned)
            
    bullets_text = "\n".join(f"- {b}" for b in summary_bullets[:6])
    if not bullets_text:
        bullets_text = "- No detailed direct context matches could be automatically formatted."
    
    fallback_text = (
        "🤖 **[Offline Fallback Mode]** I retrieved the following information from your documents:\n\n"
        f"{bullets_text}\n\n"
        "*(Note: The Gemini API is currently rate-limited or unavailable, so this is a local fallback search. "
        "Please try again in a few moments for full synthesis.)*"
    )
    return fallback_text


async def generate_fallback_answer_stream(query: str, context: str, sources: List[Dict]) -> AsyncGenerator[str, None]:
    """Stream local fallback response chunk-by-chunk to simulate streaming."""
    fallback_text = generate_fallback_rag_answer(query, context)
    chunk_size = 12
    for i in range(0, len(fallback_text), chunk_size):
        yield fallback_text[i:i+chunk_size]
        await asyncio.sleep(0.01)
    
    import json
    yield f"\n\n__SOURCES__{json.dumps(sources)}"


def generate_fallback_summary(content: str, original_name: str) -> Dict:
    """Local text-extraction summary helper."""
    lines = [line.strip() for line in content.split('\n') if line.strip()]
    key_points = []
    for line in lines:
        if line.startswith(('-', '*', '1.', '2.', '3.', '4.', '5.')):
            cleaned = re.sub(r'^[-*\d.\s]+', '', line).strip()
            if cleaned and cleaned not in key_points and len(cleaned) > 10:
                key_points.append(cleaned)
    if len(key_points) < 3:
        key_points = [line[:120] for line in lines if len(line) > 30 and not line.startswith('[Source')][:5]
        
    topics = []
    words = re.findall(r'\b[a-zA-Z]{5,}\b', content.lower())
    from collections import Counter
    common = Counter(words).most_common(20)
    stop_words = {'about', 'their', 'there', 'would', 'which', 'document', 'content', 'system', 'artificial', 'healthcare', 'medical'}
    for w, _ in common:
        if w not in stop_words and len(topics) < 3:
            topics.append(w.capitalize())
            
    summary_text = (
        f"This document '{original_name}' contains information about "
        f"{', '.join(topics) if topics else 'various topics'}. "
        f"It discusses details such as: {', '.join(key_points[:3])}."
    )
    
    return {
        "summary": summary_text,
        "key_points": key_points[:5],
        "topics": topics,
        "document_type": "Document (Local Fallback Analysis)",
        "offline": True
    }


def generate_fallback_resume_analysis(resume_text: str, job_description: Optional[str] = None) -> Dict:
    """Keyword-based local resume ATS analyzer."""
    resume_lower = resume_text.lower()
    
    skills_db = [
        "python", "react", "sql", "javascript", "typescript", "fastapi", "django", "flask",
        "docker", "kubernetes", "aws", "gcp", "git", "ci/cd", "postgres", "mongodb", "sqlite",
        "html", "css", "nodejs", "express", "machine learning", "deep learning", "nlp"
    ]
    
    found_skills = []
    for skill in skills_db:
        if re.search(r'\b' + re.escape(skill) + r'\b', resume_lower):
            found_skills.append(skill.title() if skill not in ["sql", "aws", "gcp", "ci/cd", "html", "css", "nlp"] else skill.upper())
            
    jd_skills = []
    if job_description:
        jd_lower = job_description.lower()
        for skill in skills_db:
            if re.search(r'\b' + re.escape(skill) + r'\b', jd_lower):
                jd_skills.append(skill)
                
    missing_skills = []
    if jd_skills:
        for skill in jd_skills:
            skill_name = skill.title() if skill not in ["sql", "aws", "gcp", "ci/cd", "html", "css", "nlp"] else skill.upper()
            if skill_name not in found_skills:
                missing_skills.append(skill_name)
                
    if not missing_skills:
        all_possible = ["Docker", "Kubernetes", "CI/CD", "AWS", "TypeScript"]
        missing_skills = [s for s in all_possible if s not in found_skills][:3]
        
    has_contact = "@" in resume_text or "phone" in resume_lower or re.search(r'\b\d{3}[-.\s]??\d{3}[-.\s]??\d{4}\b', resume_text)
    contact_score = 95 if has_contact else 40
    
    has_summary = "summary" in resume_lower or "profile" in resume_lower or "about" in resume_lower
    summary_score = 85 if has_summary else 50
    
    has_experience = "experience" in resume_lower or "work" in resume_lower or "employment" in resume_lower
    experience_score = 85 if has_experience else 30
    
    has_education = "education" in resume_lower or "university" in resume_lower or "college" in resume_lower or "degree" in resume_lower or "bs" in resume_lower or "ms" in resume_lower
    education_score = 80 if has_education else 45
    
    skills_score = min(40 + len(found_skills) * 10, 100)
    ats_score = int((contact_score + summary_score + experience_score + education_score + skills_score) / 5)
    
    grade = "A" if ats_score >= 90 else ("B" if ats_score >= 80 else ("C" if ats_score >= 70 else "D"))
        
    strengths = []
    if experience_score > 70:
        strengths.append("Structured experience history")
    if len(found_skills) > 4:
        strengths.append("Broad technical skillset")
    if has_contact:
        strengths.append("Contact details clearly provided")
    if not strengths:
        strengths = ["Resume has readable structure"]
        
    improvements = []
    if not has_summary:
        improvements.append("Add a professional summary section to outline your career goals")
    if len(found_skills) < 5:
        improvements.append("List more technical skills and developer tools you are familiar with")
    if not has_education:
        improvements.append("Include an education section detailing your academic credentials")
    if len(improvements) < 2:
        improvements.append("Quantify your achievements under work experience with metrics")
        
    recommended_roles = ["Developer"]
    if "python" in [s.lower() for s in found_skills]:
        recommended_roles.append("Python Engineer")
    if "react" in [s.lower() for s in found_skills]:
        recommended_roles.append("Frontend Developer")
    if len(recommended_roles) == 1:
        recommended_roles = ["Software Engineer", "Technical Specialist"]
        
    return {
        "ats_score": ats_score,
        "overall_grade": grade,
        "sections": {
            "contact_info": { "score": contact_score, "status": "good" if contact_score >= 80 else "needs_improvement", "notes": "Contact details verified" if has_contact else "Missing email or phone number" },
            "summary": { "score": summary_score, "status": "good" if summary_score >= 80 else "needs_improvement", "notes": "Summary section present" if has_summary else "Recommended to add summary" },
            "experience": { "score": experience_score, "status": "good" if experience_score >= 80 else "needs_improvement", "notes": "Professional history listed" if has_experience else "Work history section missing" },
            "education": { "score": education_score, "status": "good" if education_score >= 80 else "needs_improvement", "notes": "Education history included" if has_education else "Academic background details missing" },
            "skills": { "score": skills_score, "status": "good" if skills_score >= 80 else "needs_improvement", "notes": f"Identified {len(found_skills)} key skills" }
        },
        "found_skills": found_skills,
        "missing_skills": missing_skills,
        "strengths": strengths,
        "improvements": improvements,
        "keywords_found": found_skills[:5],
        "keywords_missing": missing_skills[:5],
        "formatting_issues": ["Bullet formatting could be modernized"] if not has_summary else [],
        "career_level": "Mid-Level" if experience_score > 60 else "Entry-Level",
        "recommended_roles": recommended_roles[:3],
        "offline": True
    }


# ══════════════════════════════════════════════════════════════════════════════
#  STREAMING ANSWER GENERATION
# ══════════════════════════════════════════════════════════════════════════════

async def generate_answer(
    query: str,
    collection_names: List[str],
    chat_history: List[Dict],
) -> AsyncGenerator[str, None]:
    """
    Streaming RAG answer generator.
    Yields text chunks as they arrive from Gemini.
    Returns sources as final JSON line.
    """
    retrieved = retrieve_context(query, collection_names)
    context, sources = build_context_string(retrieved)

    # Build conversation history string
    history_str = ""
    for msg in chat_history[-6:]:  # last 3 exchanges
        role = "User" if msg["role"] == "user" else "Assistant"
        history_str += f"{role}: {msg['content']}\n"

    system_prompt = f"""You are an intelligent AI document assistant. Answer questions accurately based ONLY on the provided document context.

RULES:
- Answer based on the document context provided
- Always cite your sources: mention "Source X, Page Y" when referencing information  
- If the answer is not in the documents, say "I couldn't find this information in the uploaded documents"
- Be concise, accurate, and helpful
- Use markdown formatting for clarity

CONVERSATION HISTORY:
{history_str}

DOCUMENT CONTEXT:
{context}

USER QUESTION: {query}

Provide a comprehensive answer with source citations:"""

    try:
        model = genai.GenerativeModel("gemini-2.0-flash")
        response = generate_content_with_retry(model, system_prompt, stream=True)

        full_text = ""
        for chunk in response:
            if chunk.text:
                full_text += chunk.text
                yield chunk.text
    except Exception as e:
        print(f"[WARN] Gemini API failed: {e}. Switching to offline fallback RAG generator...")
        async for chunk in generate_fallback_answer_stream(query, context, sources):
            yield chunk
        return

    # Yield sources as a special marker at the end
    import json
    yield f"\n\n__SOURCES__{json.dumps(sources)}"


# ══════════════════════════════════════════════════════════════════════════════
#  SUMMARIZATION
# ══════════════════════════════════════════════════════════════════════════════

def summarize_document(collection_name: str, original_name: str) -> Dict:
    """Generate summary, key points and topics for a document."""
    try:
        vs = get_vectorstore(collection_name)
        docs = vs.similarity_search("main topic summary overview", k=10)
    except Exception as e:
        print(f"[WARN] Failed to search vector store for summary: {e}. Attempting basic chunk retrieval...")
        try:
            vs = get_vectorstore(collection_name)
            res = vs.get()
            docs = []
            if res and "documents" in res:
                for doc_text, meta in zip(res["documents"], res["metadatas"]):
                    docs.append(Document(page_content=doc_text, metadata=meta))
            docs = docs[:10]
        except Exception as e_inner:
            return {"summary": f"Error retrieving document content locally: {e_inner}", "key_points": [], "topics": []}

    if not docs:
        return {"summary": "No content found.", "key_points": [], "topics": []}

    content = "\n\n".join(d.page_content for d in docs)[:8000]

    prompt = f"""Analyze this document and provide a structured analysis:

Document: {original_name}
Content: {content}

Provide your response in this exact JSON format:
{{
  "summary": "A comprehensive 3-4 sentence summary of the document",
  "key_points": ["Point 1", "Point 2", "Point 3", "Point 4", "Point 5"],
  "topics": ["Topic 1", "Topic 2", "Topic 3"],
  "document_type": "The type/category of this document"
}}"""

    try:
        model = genai.GenerativeModel("gemini-2.0-flash")
        response = generate_content_with_retry(model, prompt)
        text = response.text.strip()
    except Exception as e:
        print(f"[WARN] Gemini API summary failed: {e}. Switching to offline fallback summarizer...")
        return generate_fallback_summary(content, original_name)

    # Extract JSON from response
    import json
    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except Exception:
            pass
    return {"summary": text, "key_points": [], "topics": []}


# ══════════════════════════════════════════════════════════════════════════════
#  RESUME ANALYZER
# ══════════════════════════════════════════════════════════════════════════════

def analyze_resume(file_path: str, file_type: str, job_description: Optional[str] = None) -> Dict:
    """Full ATS-style resume analysis using Gemini."""
    docs = load_document(file_path, file_type)
    resume_text = "\n".join(d.page_content for d in docs)[:6000]

    jd_section = f"\nJOB DESCRIPTION:\n{job_description}" if job_description else ""

    prompt = f"""You are an expert ATS (Applicant Tracking System) and career coach. Analyze this resume thoroughly.
{jd_section}

RESUME:
{resume_text}

Return a comprehensive analysis in this exact JSON format:
{{
  "ats_score": 78,
  "overall_grade": "B+",
  "sections": {{
    "contact_info": {{ "score": 90, "status": "good", "notes": "..." }},
    "summary": {{ "score": 75, "status": "needs_improvement", "notes": "..." }},
    "experience": {{ "score": 85, "status": "good", "notes": "..." }},
    "education": {{ "score": 80, "status": "good", "notes": "..." }},
    "skills": {{ "score": 70, "status": "needs_improvement", "notes": "..." }}
  }},
  "found_skills": ["Python", "React", "SQL"],
  "missing_skills": ["Docker", "Kubernetes", "CI/CD"],
  "strengths": ["Strong technical background", "Quantified achievements"],
  "improvements": [
    "Add more quantified achievements with metrics",
    "Include a professional summary section",
    "Add relevant certifications"
  ],
  "keywords_found": ["machine learning", "REST API"],
  "keywords_missing": ["agile", "microservices"],
  "formatting_issues": ["Inconsistent bullet point style"],
  "career_level": "Mid-Level",
  "recommended_roles": ["Software Engineer", "Full Stack Developer"]
}}"""

    try:
        model = genai.GenerativeModel("gemini-2.0-flash")
        response = generate_content_with_retry(model, prompt)
        text = response.text.strip()
    except Exception as e:
        print(f"[WARN] Gemini API resume analysis failed: {e}. Switching to offline fallback analyzer...")
        return generate_fallback_resume_analysis(resume_text, job_description)

    import json
    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except Exception:
            pass
    return {"ats_score": 0, "error": "Could not parse resume"}
